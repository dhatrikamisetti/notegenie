import os
import json
import hashlib
import secrets
from datetime import datetime
from functools import wraps

from dotenv import load_dotenv
load_dotenv()  # Load .env file

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, jsonify, flash, send_from_directory
)
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# ── Document parsers ──────────────────────────────────────────────
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

# ── App Configuration ─────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = secrets.token_hex(32)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///notegenie.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB limit

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'md'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

# ── Configure Gemini ──────────────────────────────────────────────
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')
if GEMINI_AVAILABLE and GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.0-flash')
else:
    gemini_model = None


# ══════════════════════════════════════════════════════════════════
#  DATABASE MODELS
# ══════════════════════════════════════════════════════════════════
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    first_name = db.Column(db.String(80), default='')
    last_name = db.Column(db.String(80), default='')
    bio = db.Column(db.Text, default='')
    theme = db.Column(db.String(20), default='light')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    projects = db.relationship('Project', backref='owner', lazy=True, cascade='all, delete-orphan')
    documents = db.relationship('Document', backref='owner', lazy=True, cascade='all, delete-orphan')
    notes = db.relationship('Note', backref='owner', lazy=True, cascade='all, delete-orphan')


class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, default='')
    icon = db.Column(db.String(50), default='folder')
    color = db.Column(db.String(20), default='#5048e5')
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    documents = db.relationship('Document', backref='project', lazy=True, cascade='all, delete-orphan')


class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(300), nullable=False)
    original_filename = db.Column(db.String(300), nullable=False)
    content_text = db.Column(db.Text, default='')
    file_size = db.Column(db.Integer, default=0)
    file_type = db.Column(db.String(20), default='')
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)

    notes = db.relationship('Note', backref='document', lazy=True, cascade='all, delete-orphan')
    chat_messages = db.relationship('ChatMessage', backref='document', lazy=True, cascade='all, delete-orphan')


class Note(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    content = db.Column(db.Text, default='')
    note_type = db.Column(db.String(30), default='generated')  # generated, manual
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ChatMessage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    role = db.Column(db.String(20), nullable=False)  # 'user' or 'assistant'
    content = db.Column(db.Text, nullable=False)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


# ══════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated


def get_current_user():
    if 'user_id' in session:
        return User.query.get(session['user_id'])
    return None


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(filepath, file_type):
    """Extract text content from uploaded files."""
    text = ''
    try:
        if file_type == 'pdf' and PdfReader:
            reader = PdfReader(filepath)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n\n'
        elif file_type == 'docx' and DocxDocument:
            doc = DocxDocument(filepath)
            for para in doc.paragraphs:
                text += para.text + '\n'
        elif file_type in ('txt', 'md'):
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
    except Exception as e:
        text = f'[Error extracting text: {str(e)}]'
    return text


def format_file_size(size_bytes):
    """Format bytes to human readable size."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"


def ai_chat(document_text, user_message, chat_history=None):
    """Send a message to Gemini AI with document context."""
    if not gemini_model:
        return "⚠️ AI is not configured. Please set the GEMINI_API_KEY environment variable to enable AI features.\n\nTo set it, run:\n```\nset GEMINI_API_KEY=your_api_key_here\n```\nThen restart the server."

    # Build context
    system_prompt = """You are NoteGenie AI, a brilliant research assistant. You help users understand, 
analyze, and extract insights from their documents. 

Rules:
- Always reference specific parts of the document when answering
- Provide clear, structured answers with citations when possible
- If asked to summarize, provide comprehensive but concise summaries
- Use markdown formatting in your responses
- If the document doesn't contain relevant information, say so honestly
- Be helpful, accurate, and insightful"""

    # Build conversation
    context = f"DOCUMENT CONTENT:\n{document_text[:30000]}\n\n"
    
    history_text = ""
    if chat_history:
        for msg in chat_history[-10:]:  # Last 10 messages for context
            history_text += f"{msg['role'].upper()}: {msg['content']}\n"

    full_prompt = f"{system_prompt}\n\n{context}\n{history_text}\nUSER: {user_message}\n\nASSISTANT:"

    try:
        response = gemini_model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        return f"Sorry, I encountered an error: {str(e)}"


def ai_summarize(document_text):
    """Generate a summary of the document."""
    if not gemini_model:
        return {
            'title': 'AI Not Configured',
            'summary': 'Please set the GEMINI_API_KEY environment variable to enable AI summarization.',
            'key_points': []
        }

    prompt = f"""Analyze the following document and provide a structured summary.

Return your response in this exact JSON format:
{{
    "title": "A descriptive title for the key concept",
    "summary": "A comprehensive 2-3 paragraph summary",
    "key_points": ["Key point 1", "Key point 2", "Key point 3", "Key point 4", "Key point 5"]
}}

DOCUMENT:
{document_text[:30000]}
"""
    try:
        response = gemini_model.generate_content(prompt)
        text = response.text.strip()
        # Try to parse JSON from response
        if '```json' in text:
            text = text.split('```json')[1].split('```')[0].strip()
        elif '```' in text:
            text = text.split('```')[1].split('```')[0].strip()
        return json.loads(text)
    except Exception as e:
        return {
            'title': 'Document Summary',
            'summary': f'Error generating summary: {str(e)}',
            'key_points': []
        }


# ══════════════════════════════════════════════════════════════════
#  ROUTES — PUBLIC PAGES
# ══════════════════════════════════════════════════════════════════
@app.route('/')
def landing():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template('landing.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        
        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            session['user_name'] = f"{user.first_name} {user.last_name}".strip() or email
            return redirect(url_for('dashboard'))
        
        flash('Invalid email or password.', 'error')
    
    return render_template('login.html')


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()

        if User.query.filter_by(email=email).first():
            flash('An account with this email already exists.', 'error')
            return render_template('login.html', show_signup=True)

        if len(password) < 6:
            flash('Password must be at least 6 characters.', 'error')
            return render_template('login.html', show_signup=True)

        user = User(
            email=email,
            password_hash=generate_password_hash(password),
            first_name=first_name,
            last_name=last_name
        )
        db.session.add(user)
        db.session.commit()

        # Create a default project
        default_project = Project(
            name='My First Project',
            description='Your default project workspace',
            icon='science',
            user_id=user.id
        )
        db.session.add(default_project)
        db.session.commit()

        session['user_id'] = user.id
        session['user_name'] = f"{first_name} {last_name}".strip() or email
        return redirect(url_for('dashboard'))

    return render_template('login.html', show_signup=True)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('landing'))


# ══════════════════════════════════════════════════════════════════
#  ROUTES — DASHBOARD
# ══════════════════════════════════════════════════════════════════
@app.route('/dashboard')
@login_required
def dashboard():
    user = get_current_user()
    projects = Project.query.filter_by(user_id=user.id).order_by(Project.updated_at.desc()).all()
    documents = Document.query.filter_by(user_id=user.id).order_by(Document.uploaded_at.desc()).all()
    notes = Note.query.filter_by(user_id=user.id).order_by(Note.created_at.desc()).all()
    
    stats = {
        'total_documents': len(documents),
        'active_projects': len(projects),
        'total_notes': len(notes)
    }
    
    return render_template('dashboard.html', 
                         user=user, 
                         projects=projects, 
                         documents=documents,
                         notes=notes,
                         stats=stats,
                         format_file_size=format_file_size)


# ══════════════════════════════════════════════════════════════════
#  ROUTES — PROJECTS
# ══════════════════════════════════════════════════════════════════
@app.route('/api/projects', methods=['POST'])
@login_required
def create_project():
    user = get_current_user()
    data = request.get_json()
    
    project = Project(
        name=data.get('name', 'Untitled Project'),
        description=data.get('description', ''),
        icon=data.get('icon', 'folder'),
        color=data.get('color', '#5048e5'),
        user_id=user.id
    )
    db.session.add(project)
    db.session.commit()
    
    return jsonify({
        'id': project.id,
        'name': project.name,
        'description': project.description,
        'message': 'Project created successfully'
    })


@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
@login_required
def delete_project(project_id):
    user = get_current_user()
    project = Project.query.filter_by(id=project_id, user_id=user.id).first_or_404()
    db.session.delete(project)
    db.session.commit()
    return jsonify({'message': 'Project deleted'})


# ══════════════════════════════════════════════════════════════════
#  ROUTES — DOCUMENT UPLOAD & MANAGEMENT
# ══════════════════════════════════════════════════════════════════
@app.route('/api/upload', methods=['POST'])
@login_required
def upload_document():
    user = get_current_user()
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': f'File type not allowed. Supported: {", ".join(ALLOWED_EXTENSIONS)}'}), 400
    
    project_id = request.form.get('project_id', type=int)
    
    # Save file
    original_filename = file.filename
    ext = original_filename.rsplit('.', 1)[1].lower()
    safe_name = secure_filename(f"{user.id}_{secrets.token_hex(8)}.{ext}")
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_name)
    file.save(filepath)
    
    # Get file size
    file_size = os.path.getsize(filepath)
    
    # Extract text
    content_text = extract_text(filepath, ext)
    
    # Save to database
    doc = Document(
        filename=safe_name,
        original_filename=original_filename,
        content_text=content_text,
        file_size=file_size,
        file_type=ext,
        project_id=project_id,
        user_id=user.id
    )
    db.session.add(doc)
    db.session.commit()
    
    return jsonify({
        'id': doc.id,
        'filename': original_filename,
        'file_size': format_file_size(file_size),
        'file_type': ext,
        'message': 'Document uploaded successfully',
        'has_content': bool(content_text.strip())
    })


@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
@login_required
def delete_document(doc_id):
    user = get_current_user()
    doc = Document.query.filter_by(id=doc_id, user_id=user.id).first_or_404()
    
    # Delete file from disk
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], doc.filename)
    if os.path.exists(filepath):
        os.remove(filepath)
    
    db.session.delete(doc)
    db.session.commit()
    return jsonify({'message': 'Document deleted'})


# ══════════════════════════════════════════════════════════════════
#  ROUTES — WORKSPACE
# ══════════════════════════════════════════════════════════════════
@app.route('/workspace')
@app.route('/workspace/<int:doc_id>')
@login_required
def workspace(doc_id=None):
    user = get_current_user()
    documents = Document.query.filter_by(user_id=user.id).order_by(Document.uploaded_at.desc()).all()
    projects = Project.query.filter_by(user_id=user.id).all()
    
    current_doc = None
    chat_messages = []
    doc_notes = []
    
    if doc_id:
        current_doc = Document.query.filter_by(id=doc_id, user_id=user.id).first()
        if current_doc:
            chat_messages = ChatMessage.query.filter_by(
                document_id=doc_id, user_id=user.id
            ).order_by(ChatMessage.created_at.asc()).all()
            doc_notes = Note.query.filter_by(
                document_id=doc_id, user_id=user.id
            ).order_by(Note.created_at.desc()).all()
    
    return render_template('workspace.html',
                         user=user,
                         documents=documents,
                         projects=projects,
                         current_doc=current_doc,
                         chat_messages=chat_messages,
                         doc_notes=doc_notes,
                         format_file_size=format_file_size)


# ══════════════════════════════════════════════════════════════════
#  ROUTES — AI CHAT
# ══════════════════════════════════════════════════════════════════
@app.route('/api/chat', methods=['POST'])
@login_required
def chat():
    user = get_current_user()
    data = request.get_json()
    
    doc_id = data.get('document_id')
    message = data.get('message', '').strip()
    
    if not doc_id or not message:
        return jsonify({'error': 'Document ID and message are required'}), 400
    
    doc = Document.query.filter_by(id=doc_id, user_id=user.id).first()
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    
    # Save user message
    user_msg = ChatMessage(
        role='user',
        content=message,
        document_id=doc_id,
        user_id=user.id
    )
    db.session.add(user_msg)
    db.session.commit()
    
    # Get chat history
    history = ChatMessage.query.filter_by(
        document_id=doc_id, user_id=user.id
    ).order_by(ChatMessage.created_at.asc()).all()
    
    chat_history = [{'role': m.role, 'content': m.content} for m in history]
    
    # Get AI response
    ai_response = ai_chat(doc.content_text, message, chat_history)
    
    # Save assistant message
    assistant_msg = ChatMessage(
        role='assistant',
        content=ai_response,
        document_id=doc_id,
        user_id=user.id
    )
    db.session.add(assistant_msg)
    db.session.commit()
    
    return jsonify({
        'response': ai_response,
        'message_id': assistant_msg.id
    })


@app.route('/api/summarize', methods=['POST'])
@login_required
def summarize():
    user = get_current_user()
    data = request.get_json()
    doc_id = data.get('document_id')
    
    doc = Document.query.filter_by(id=doc_id, user_id=user.id).first()
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    
    summary = ai_summarize(doc.content_text)
    
    # Save as a note
    note = Note(
        title=summary.get('title', 'Document Summary'),
        content=json.dumps(summary),
        note_type='generated',
        document_id=doc_id,
        user_id=user.id
    )
    db.session.add(note)
    db.session.commit()
    
    return jsonify({
        'note_id': note.id,
        'summary': summary
    })


# ══════════════════════════════════════════════════════════════════
#  ROUTES — NOTES
# ══════════════════════════════════════════════════════════════════
@app.route('/api/notes', methods=['POST'])
@login_required
def save_note():
    user = get_current_user()
    data = request.get_json()
    
    note = Note(
        title=data.get('title', 'Untitled Note'),
        content=data.get('content', ''),
        note_type=data.get('note_type', 'manual'),
        document_id=data.get('document_id'),
        user_id=user.id
    )
    db.session.add(note)
    db.session.commit()
    
    return jsonify({
        'id': note.id,
        'title': note.title,
        'message': 'Note saved'
    })


@app.route('/api/notes/<int:note_id>', methods=['DELETE'])
@login_required
def delete_note(note_id):
    user = get_current_user()
    note = Note.query.filter_by(id=note_id, user_id=user.id).first_or_404()
    db.session.delete(note)
    db.session.commit()
    return jsonify({'message': 'Note deleted'})


# ══════════════════════════════════════════════════════════════════
#  ROUTES — SETTINGS
# ══════════════════════════════════════════════════════════════════
@app.route('/settings')
@login_required
def settings():
    user = get_current_user()
    return render_template('settings.html', user=user)


@app.route('/api/settings', methods=['POST'])
@login_required
def update_settings():
    user = get_current_user()
    data = request.get_json()
    
    if 'first_name' in data:
        user.first_name = data['first_name']
    if 'last_name' in data:
        user.last_name = data['last_name']
    if 'bio' in data:
        user.bio = data['bio']
    if 'theme' in data:
        user.theme = data['theme']
    
    db.session.commit()
    session['user_name'] = f"{user.first_name} {user.last_name}".strip() or user.email
    
    return jsonify({'message': 'Settings updated successfully'})


@app.route('/api/settings/password', methods=['POST'])
@login_required
def update_password():
    user = get_current_user()
    data = request.get_json()
    
    current = data.get('current_password', '')
    new = data.get('new_password', '')
    
    if not check_password_hash(user.password_hash, current):
        return jsonify({'error': 'Current password is incorrect'}), 400
    
    if len(new) < 6:
        return jsonify({'error': 'New password must be at least 6 characters'}), 400
    
    user.password_hash = generate_password_hash(new)
    db.session.commit()
    
    return jsonify({'message': 'Password updated successfully'})


# ══════════════════════════════════════════════════════════════════
#  INITIALIZE & RUN
# ══════════════════════════════════════════════════════════════════
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    print("\n" + "="*60)
    print("  🧠 NoteGenie — AI-Powered Research Assistant")
    print("="*60)
    if gemini_model:
        print("  ✅ Gemini AI is configured and ready")
    else:
        print("  ⚠️  Set GEMINI_API_KEY env variable for AI features")
    print(f"  📂 Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f"  🌐 Open: http://127.0.0.1:5000")
    print("="*60 + "\n")
    app.run(debug=True, port=5000)
